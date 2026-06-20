#!/usr/bin/env python3
"""Production fix pass for Baren Sump trilogy manuscripts."""

import re
import shutil
import zipfile
from pathlib import Path

UPLOADS = Path("/home/ubuntu/.cursor/projects/workspace/uploads/extracted")
OUT = Path("/workspace/production/books")

# In-run fixes (entire string inside one <w:t> node)
IN_RUN_FIXES = [
    ("triedTHELASTPRESIDENT", "tried THELASTPRESIDENT"),  # split further below if needed
]

# When t1 ends with this suffix and t2 starts with uppercase, insert space.
NEEDS_SPACE_SUFFIXES = (
    "read",
    "said",
    "called",
    "written",
    "spelled it",
    "titled",
    "labeled",
    "beneath",
    "tried",
    " searched",
    "for",
    "used",
    "word",
    "words",
    "to",
    "shouted",
    "circled",
    "underlined",
    "crossed out",
    "asked",
    "called itself an",
    "called it",
    "phrase",
    "received",
    "At",
    "an",
    "the",
    "it",
    "file",
    "searched",
)

WT = re.compile(
    r"(<w:t(?: xml:space=\"preserve\")?>)([^<]*?)(</w:t>)",
    re.S,
)


def needs_space_between(t1: str, t2: str) -> bool:
    if not t1 or not t2:
        return False
    if t2[0] == " " or t1[-1] == " ":
        return False
    if t1[-1] == "-":
        return False
    # Adjacent runs glued without a space (bold/italic splits)
    if t1[-1].isalnum() and t2[0].isalnum():
        return True
    if any(t1.endswith(s) for s in NEEDS_SPACE_SUFFIXES) and t2[0].isupper():
        return True
    return False


def fix_wt_content(text: str) -> str:
    for old, new in IN_RUN_FIXES:
        text = text.replace(old, new)
    # split glued ALL CAPS after common verbs inside same run
    text = re.sub(
        r"\b(tried|read|said|used|At)([A-Z]{2,})",
        r"\1 \2",
        text,
    )
    text = re.sub(
        r"\b(circled)(laughed)\b",
        r"\1 \2",
        text,
    )
    text = re.sub(
        r"\b(underlined)(function)\b",
        r"\1 \2",
        text,
    )
    text = re.sub(
        r"\b(crossed out)(may)\b",
        r"\1 \2",
        text,
    )
    text = re.sub(
        r"\b(asked)(well)\b",
        r"\1 \2",
        text,
    )
    text = re.sub(
        r"\b(and)(stabilization)\b",
        r"\1 \2",
        text,
    )
    text = re.sub(
        r"\b(out)(may)\b",
        r"\1 \2",
        text,
    )
    text = re.sub(
        r"THELASTPRESIDENT",
        "THE LAST PRESIDENT",
        text,
    )
    return text


def fix_paragraph_xml(p_xml: str) -> tuple[str, int]:
    changes = 0

    def fix_text(m: re.Match) -> str:
        nonlocal changes
        open_tag, content, close_tag = m.group(1), m.group(2), m.group(3)
        new_content = fix_wt_content(content)
        if new_content != content:
            changes += 1
            content = new_content
        return open_tag + content + close_tag

    p_xml = WT.sub(fix_text, p_xml)

    # Fix missing spaces across adjacent runs in same paragraph
    parts = WT.split(p_xml)
    # parts: [before, open1, text1, close1, between, open2, text2, close2, ...]
    if len(parts) >= 4:
        rebuilt = [parts[0]]
        i = 1
        prev_text = None
        while i + 2 < len(parts):
            open_tag, text, close_tag = parts[i], parts[i + 1], parts[i + 2]
            if prev_text is not None and needs_space_between(prev_text, text):
                if not text.startswith(" "):
                    text = " " + text
                    open_tag = open_tag.replace("<w:t>", '<w:t xml:space="preserve">').replace(
                        '<w:t xml:space="preserve">',
                        '<w:t xml:space="preserve">',
                        1,
                    )
                    if 'xml:space="preserve"' not in open_tag:
                        open_tag = '<w:t xml:space="preserve">'
                    changes += 1
            rebuilt.extend([open_tag, text, close_tag])
            prev_text = text
            i += 3
            if i < len(parts):
                rebuilt.append(parts[i])
                i += 1
        p_xml = "".join(rebuilt)

    return p_xml, changes


def fix_document_xml(xml: str) -> tuple[str, int]:
    total = 0
    paragraphs = re.split(r"(</w:p>)", xml)
    out = []
    for i in range(0, len(paragraphs), 2):
        chunk = paragraphs[i]
        closer = paragraphs[i + 1] if i + 1 < len(paragraphs) else ""
        if "<w:p" in chunk:
            fixed, n = fix_paragraph_xml(chunk + closer)
            total += n
            out.append(fixed)
        else:
            out.append(chunk + closer)
    return "".join(out), total


def patch_docx(src: Path, dst: Path) -> int:
    changes = 0
    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                xml = data.decode("utf-8")
                xml, n = fix_document_xml(xml)
                changes += n
                data = xml.encode("utf-8")
            zout.writestr(item, data)
    return changes


def style_interlude_book_one(docx_path: Path) -> None:
    from docx import Document

    doc = Document(str(docx_path))
    for para in doc.paragraphs:
        t = para.text.strip()
        if t == "INTERLUDE":
            para.style = doc.styles["Heading 1"]
        elif t == "The Room That Measured Him":
            para.style = doc.styles["Heading 2"]
    doc.save(str(docx_path))


def add_chapter_thirty_subtitle(docx_path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_BREAK

    doc = Document(str(docx_path))
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "CHAPTER THIRTY" and para.style.name.startswith("Heading"):
            nxt = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
            if nxt and not nxt.style.name.startswith("Heading"):
                new_para = nxt.insert_paragraph_before("They Carried Moral Force")
                new_para.style = doc.styles["Heading 2"]
                break
    doc.save(str(docx_path))


def remove_duplicate_black_path_title(docx_path: Path) -> None:
    from docx import Document

    doc = Document(str(docx_path))
    plain_titles = []
    for para in doc.paragraphs[:20]:
        if para.text.strip() == "The Black Path" and not para.style.name.startswith("Heading"):
            plain_titles.append(para)
    if len(plain_titles) >= 2:
        p = plain_titles[1]._element
        p.getparent().remove(p)
    doc.save(str(docx_path))


PACKAGE_MAP = {
    "book-one": {
        "source_dir": UPLOADS / "Baren_Sump_BOOK_ONE_FINAL_PACKAGE_3f64",
        "reader": "Baren_Sump_BOOK_ONE_The_Last_President_FINAL_READER_COPY.docx",
        "master": "Baren_Sump_BOOK_ONE_The_Last_President_AUTHOR_PRODUCTION_MASTER.docx",
        "reader_out": "Baren_Sump_BOOK_ONE_The_Last_President_FINAL_READER_COPY.docx",
        "master_out": "Baren_Sump_BOOK_ONE_The_Last_President_AUTHOR_PRODUCTION_MASTER.docx",
        "extras": [
            "Baren_Sump_BOOK_ONE_Art_Brief.docx",
            "Baren_Sump_BOOK_ONE_Final_Package_Log.docx",
        ],
        "post": "book-one",
    },
    "book-two": {
        "source_dir": UPLOADS / "Baren_Sump_BOOK_TWO_COMPRESSION_VOICE_PASS_PACKAGE_90e0",
        "reader": "Baren_Sump_BOOK_TWO_Children_of_Tomorrow_COMPRESSION_VOICE_PASS.docx",
        "master": "Baren_Sump_BOOK_TWO_Children_of_Tomorrow_AUTHOR_MASTER_COMPRESSION_PASS.docx",
        "reader_out": "Baren_Sump_BOOK_TWO_Children_of_Tomorrow_FINAL_READER_COPY.docx",
        "master_out": "Baren_Sump_BOOK_TWO_Children_of_Tomorrow_AUTHOR_PRODUCTION_MASTER.docx",
        "extras": ["Baren_Sump_BOOK_TWO_Compression_Voice_Pass_Log.docx"],
        "post": "book-two",
    },
    "book-three": {
        "source_dir": UPLOADS / "Baren_Sump_BOOK_THREE_COMPRESSION_VOICE_PASS_PACKAGE_9b4b",
        "reader": "Baren_Sump_BOOK_THREE_The_Black_Path_COMPRESSION_VOICE_PASS.docx",
        "master": "Baren_Sump_BOOK_THREE_The_Black_Path_AUTHOR_MASTER_COMPRESSION_PASS.docx",
        "reader_out": "Baren_Sump_BOOK_THREE_The_Black_Path_FINAL_READER_COPY.docx",
        "master_out": "Baren_Sump_BOOK_THREE_The_Black_Path_AUTHOR_PRODUCTION_MASTER.docx",
        "extras": ["Baren_Sump_BOOK_THREE_Compression_Voice_Pass_Log.docx"],
        "post": "book-three",
    },
}


def write_production_log(book_dir: Path, book_name: str, changes: int, notes: list[str]) -> None:
    (book_dir / "PRODUCTION_LOG.txt").write_text(
        f"{book_name} — Final Production Pass\nDate: 2026-06-20\nRun-boundary / text fixes: {changes}\n\n"
        + "\n".join(f"- {n}" for n in notes)
        + "\n",
        encoding="utf-8",
    )


def verify_book_one(path: Path) -> list[str]:
    import zipfile

    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    joined = re.sub(r"<[^>]+>", "", xml)
    bad = [
        "writtenChildren",
        "calledThe",
        "circledlaughed",
        "andstabilization",
        "askedwell",
        "readTHE",
        "AtBAREN",
        "triedTHELASTPRESIDENT",
    ]
    return [b for b in bad if b in joined]


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for book_key, cfg in PACKAGE_MAP.items():
        book_dir = OUT / book_key
        book_dir.mkdir(parents=True, exist_ok=True)
        total = 0
        notes = []
        for role in ("reader", "master"):
            src = cfg["source_dir"] / cfg[role]
            dst = book_dir / cfg[f"{role}_out"]
            tmp = book_dir / f"_tmp_{cfg[f'{role}_out']}"
            n = patch_docx(src, tmp)
            total += n
            shutil.move(tmp, dst)
            print(f"  {cfg[f'{role}_out']}: {n} fixes")

        reader = book_dir / cfg["reader_out"]
        master = book_dir / cfg["master_out"]

        if cfg["post"] == "book-one":
            style_interlude_book_one(reader)
            style_interlude_book_one(master)
            notes.append("Applied Heading styles to Book One INTERLUDE")
            remaining = verify_book_one(reader)
            if remaining:
                notes.append(f"WARNING remaining tokens: {remaining}")
            else:
                notes.append("Verified: all known Book One spacing defects resolved")

        if cfg["post"] == "book-two":
            add_chapter_thirty_subtitle(reader)
            add_chapter_thirty_subtitle(master)
            notes.append("Added Chapter Thirty subtitle: They Carried Moral Force")
            notes.append("Renamed to FINAL_READER_COPY / AUTHOR_PRODUCTION_MASTER")

        if cfg["post"] == "book-three":
            remove_duplicate_black_path_title(reader)
            remove_duplicate_black_path_title(master)
            notes.append("Removed duplicate title-page line")
            notes.append("Renamed to FINAL_READER_COPY / AUTHOR_PRODUCTION_MASTER")

        for extra in cfg.get("extras", []):
            src = cfg["source_dir"] / extra
            if src.exists():
                shutil.copy2(src, book_dir / extra)

        write_production_log(book_dir, book_key, total, notes)
        print(f"OK {book_key}")


if __name__ == "__main__":
    main()
