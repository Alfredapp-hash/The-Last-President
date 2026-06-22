#!/usr/bin/env python3
"""Full-manuscript integrity scan for sale-ready release."""

from __future__ import annotations

import argparse
import re
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from docx import Document


ROOT = Path("/workspace")
BOOKS_DIR = ROOT / "production" / "books"
EXPORTS_DIR = ROOT / "production" / "publication-prep" / "exports"
OUT_DIR = ROOT / "production" / "publication-prep" / "sale-ready"

HARD_FLAG_PATTERNS = [
    (r"\[ILLUSTRATION PLACEMENT\]", "illustration placement"),
    (r"\bAUTHOR(?:'S)?\s+NOTE\b", "author note"),
    (r"\bEDITOR(?:'S)?\s+NOTE\b", "editor note"),
    (r"\bPRODUCTION\s+NOTE\b", "production note"),
    (r"\bTODO\b", "TODO"),
    (r"\bFIXME\b", "FIXME"),
    (r"\bTK\b", "TK"),
    (r"\bTBD\b", "TBD"),
    (r"\bPLACEHOLDER\b", "PLACEHOLDER"),
    (r"\[NOTE[:\s]", "bracketed NOTE"),
    (r"\[EDIT[:\s]", "bracketed EDIT"),
    (r"<!--", "HTML comment"),
    (r"\bLorem ipsum\b", "lorem ipsum"),
    (r"\bXXX\b", "XXX marker"),
    (r"\?\?\?", "??? marker"),
    (r"\[ART\b", "ART bracket"),
    (r"\bART\s+BRIEF\b", "ART BRIEF"),
    (r"\bQUERY:\b", "QUERY"),
    (r"\bSTET\b", "STET"),
    (r"\[CHAPTER\s+STUB\]", "chapter stub"),
    (r"\bNEEDS\s+(?:REWRITE|EDIT|WORK)\b", "needs work"),
    (r"\bCHECK\s+WITH\s+AUTHOR\b", "check with author"),
    (r"\[internal\]", "internal bracket"),
    (r"You.re around Chapter", "chapter-progress author note"),
    (r"I.d estimate \d+[-–]\d+ chapters left", "chapter-count author note"),
    (r"remaining arc should probably be", "outline author note"),
    (r"late Act II\s*/\s*early Act III", "structure author note"),
    (r"For this satirical thriller, I.d treat", "genre author note"),
    (r"\bAs an AI\b", "AI assistant note"),
    (r"\bChatGPT\b", "AI assistant note"),
    (r"\bwrite this later\b", "draft instruction"),
    (r"\bINSERT HERE\b", "insert instruction"),
    (r"\bDRAFT ONLY\b", "draft-only marker"),
]

ALLOWLIST_LINE_PATTERNS = [
    re.compile(
        r"PREDICTION → USER ACTION → EVENT ALTERATION → ATTENTION → REINFORCED PREDICTION → USER DEPENDENCE",
        re.I,
    ),
]

SPACING_DEFECTS = [
    "writtenChildren",
    "calledThe",
    "circledlaughed",
    "andstabilization",
    "askedwell",
    "readTHE",
    "AtBAREN",
    "triedTHELASTPRESIDENT",
    "underlinedfunction",
    "outmay",
]

EPUB_PATHS = {
    "The Last President": EXPORTS_DIR / "epub/baren-sump-book-1-the-last-president_v1.epub",
    "Children of Tomorrow": EXPORTS_DIR / "epub/baren-sump-book-2-children-of-tomorrow_v1.epub",
    "The Black Path": EXPORTS_DIR / "epub/baren-sump-book-3-the-black-path_v1.epub",
}


@dataclass
class BookConfig:
    key: str
    title: str
    reader_docx: str
    end_marker: str


BOOKS = [
    BookConfig(
        key="book-one",
        title="The Last President",
        reader_docx="Baren_Sump_BOOK_ONE_The_Last_President_FINAL_READER_COPY.docx",
        end_marker="END OF BOOK ONE",
    ),
    BookConfig(
        key="book-two",
        title="Children of Tomorrow",
        reader_docx="Baren_Sump_BOOK_TWO_Children_of_Tomorrow_FINAL_READER_COPY.docx",
        end_marker="END OF BOOK TWO",
    ),
    BookConfig(
        key="book-three",
        title="The Black Path",
        reader_docx="Baren_Sump_BOOK_THREE_The_Black_Path_FINAL_READER_COPY.docx",
        end_marker="END OF BOOK THREE",
    ),
]


def words_count(text: str) -> int:
    return len(re.findall(r"\b[\w'’-]+\b", text))


def is_allowlisted(text: str) -> bool:
    return any(pat.search(text) for pat in ALLOWLIST_LINE_PATTERNS)


def scan_docx(path: Path, title: str, end_marker: str) -> tuple[dict, list[dict]]:
    doc = Document(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    stats = {
        "paragraphs": sum(1 for p in doc.paragraphs if p.text.strip()),
        "words": words_count(full_text),
        "has_end_marker": end_marker in full_text,
    }
    issues: list[dict] = []
    current_h1 = ""
    for i, p in enumerate(doc.paragraphs, 1):
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style else ""
        if style == "Heading 1":
            current_h1 = text
        if is_allowlisted(text):
            continue
        for pat, label in HARD_FLAG_PATTERNS:
            if re.search(pat, text, re.I):
                issues.append(
                    {
                        "book": title,
                        "line": i,
                        "section": current_h1,
                        "label": label,
                        "context": text[:240],
                    }
                )
                break
        for defect in SPACING_DEFECTS:
            if defect in text:
                issues.append(
                    {
                        "book": title,
                        "line": i,
                        "section": current_h1,
                        "label": "spacing defect",
                        "context": text[:240],
                    }
                )
                break
    if not stats["has_end_marker"]:
        issues.append(
            {
                "book": title,
                "line": None,
                "section": "",
                "label": "missing end marker",
                "context": end_marker,
            }
        )
    return stats, issues


def scan_epub(path: Path, title: str) -> list[dict]:
    if not path.exists():
        return [
            {
                "book": title,
                "line": None,
                "section": "EPUB",
                "label": "missing export",
                "context": str(path),
            }
        ]
    issues: list[dict] = []
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            if not (name.endswith(".xhtml") or name.endswith(".html")):
                continue
            text = zf.read(name).decode("utf-8", errors="ignore")
            for pat, label in HARD_FLAG_PATTERNS:
                if re.search(pat, text, re.I):
                    issues.append(
                        {
                            "book": title,
                            "line": None,
                            "section": f"EPUB:{name}",
                            "label": label,
                            "context": "pattern found in export",
                        }
                    )
                    break
    return issues


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--out",
        default=str(OUT_DIR / "FULL_MANUSCRIPT_SCAN.md"),
        help="Report output path",
    )
    args = parser.parse_args()

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    all_issues: list[dict] = []
    lines = [
        "# Final Full Manuscript Scan",
        "",
        f"Date: {now}",
        "",
        "Scope: every paragraph in all three final reader copies, plus EPUB export verification.",
        "",
    ]

    lines.append("## Manuscript scan")
    lines.append("")
    lines.append("| Book | Words | Paragraphs | End marker | Hard flags |")
    lines.append("|------|------:|-----------:|:----------:|-----------:|")

    for cfg in BOOKS:
        path = BOOKS_DIR / cfg.key / cfg.reader_docx
        stats, issues = scan_docx(path, cfg.title, cfg.end_marker)
        epub_issues = scan_epub(EPUB_PATHS[cfg.title], cfg.title)
        book_issues = issues + epub_issues
        all_issues.extend(book_issues)
        lines.append(
            f"| {cfg.title} | {stats['words']:,} | {stats['paragraphs']:,} | "
            f"{'yes' if stats['has_end_marker'] else 'no'} | {len(book_issues)} |"
        )

    lines.extend(["", "## Hard-flag details", ""])
    if all_issues:
        for issue in all_issues:
            loc = f"L{issue['line']}" if issue["line"] else issue["section"]
            lines.append(
                f"- **{issue['book']}** [{issue['label']}] {loc}: {issue['context']}"
            )
    else:
        lines.append("- No hard flags detected.")

    lines.extend(
        [
            "",
            "## In-story terms reviewed (not flagged)",
            "",
            "- `draft` / `Draft` used in legal and editorial dialogue",
            "- In-story flow notation in Book One Chapter Fourteen",
            "",
        ]
    )

    status = "PASS" if not all_issues else "FLAGGED"
    lines.insert(4, f"Overall: **{status}**")
    lines.insert(5, "")

    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(f"Wrote {out_path}")
    print(f"Overall: {status} ({len(all_issues)} hard flags)")
    if all_issues:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
