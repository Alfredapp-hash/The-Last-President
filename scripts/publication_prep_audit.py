#!/usr/bin/env python3
"""Final publication-prep QA audit for trilogy manuscripts."""

from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from docx import Document


ROOT = Path("/workspace")
BOOKS_DIR = ROOT / "production" / "books"
OUT_DIR = ROOT / "production" / "publication-prep" / "audit"

FLAG_PATTERNS = [
    r"\[ILLUSTRATION PLACEMENT\]",
    r"\bTK\b",
    r"\bTODO\b",
    r"\bFIXME\b",
    r"\bPLACEHOLDER\b",
]


@dataclass
class BookConfig:
    key: str
    title: str
    reader_docx: str
    master_docx: str
    end_marker: str
    start_chapter: int
    end_chapter: int
    has_prologue: bool = False
    has_interlude: bool = False


BOOKS = [
    BookConfig(
        key="book-one",
        title="The Last President",
        reader_docx="Baren_Sump_BOOK_ONE_The_Last_President_FINAL_READER_COPY.docx",
        master_docx="Baren_Sump_BOOK_ONE_The_Last_President_AUTHOR_PRODUCTION_MASTER.docx",
        end_marker="END OF BOOK ONE",
        start_chapter=1,
        end_chapter=24,
        has_prologue=True,
        has_interlude=True,
    ),
    BookConfig(
        key="book-two",
        title="Children of Tomorrow",
        reader_docx="Baren_Sump_BOOK_TWO_Children_of_Tomorrow_FINAL_READER_COPY.docx",
        master_docx="Baren_Sump_BOOK_TWO_Children_of_Tomorrow_AUTHOR_PRODUCTION_MASTER.docx",
        end_marker="END OF BOOK TWO",
        start_chapter=25,
        end_chapter=54,
        has_interlude=True,
    ),
    BookConfig(
        key="book-three",
        title="The Black Path",
        reader_docx="Baren_Sump_BOOK_THREE_The_Black_Path_FINAL_READER_COPY.docx",
        master_docx="Baren_Sump_BOOK_THREE_The_Black_Path_AUTHOR_PRODUCTION_MASTER.docx",
        end_marker="END OF BOOK THREE",
        start_chapter=55,
        end_chapter=84,
    ),
]


NUM_WORDS = {
    1: "ONE",
    2: "TWO",
    3: "THREE",
    4: "FOUR",
    5: "FIVE",
    6: "SIX",
    7: "SEVEN",
    8: "EIGHT",
    9: "NINE",
    10: "TEN",
    11: "ELEVEN",
    12: "TWELVE",
    13: "THIRTEEN",
    14: "FOURTEEN",
    15: "FIFTEEN",
    16: "SIXTEEN",
    17: "SEVENTEEN",
    18: "EIGHTEEN",
    19: "NINETEEN",
    20: "TWENTY",
    30: "THIRTY",
    40: "FORTY",
    50: "FIFTY",
    60: "SIXTY",
    70: "SEVENTY",
    80: "EIGHTY",
    90: "NINETY",
}


def num_to_words(n: int) -> str:
    if n in NUM_WORDS:
        return NUM_WORDS[n]
    if n < 100:
        tens = (n // 10) * 10
        ones = n % 10
        return f"{NUM_WORDS[tens]}-{NUM_WORDS[ones]}"
    raise ValueError(f"Unsupported chapter number: {n}")


def words_count(text: str) -> int:
    return len(re.findall(r"\b[\w'’-]+\b", text))


def text_from_doc(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def expected_headings(cfg: BookConfig) -> list[str]:
    out: list[str] = []
    if cfg.has_prologue:
        out.append("PROLOGUE")
    if cfg.has_interlude:
        out.append("INTERLUDE")
    for n in range(cfg.start_chapter, cfg.end_chapter + 1):
        out.append(f"CHAPTER {num_to_words(n)}")
    out.append(cfg.end_marker)
    return out


def heading_lines(doc: Document) -> list[tuple[int, str, str]]:
    rows = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ""
        if text and style.startswith("Heading"):
            rows.append((i, style, text))
    return rows


def audit_book(cfg: BookConfig) -> tuple[str, dict[str, int], list[str]]:
    bdir = BOOKS_DIR / cfg.key
    reader_path = bdir / cfg.reader_docx
    master_path = bdir / cfg.master_docx
    reader_doc = Document(str(reader_path))
    master_doc = Document(str(master_path))

    reader_text = text_from_doc(reader_doc)
    master_text = text_from_doc(master_doc)
    reader_headings = heading_lines(reader_doc)
    expected = expected_headings(cfg)
    h1_actual = [t for _, style, t in reader_headings if style == "Heading 1"]
    missing = [h for h in expected if h not in h1_actual]

    # Check subtitle presence: each Heading 1 chapter/prologue/interlude should be followed by Heading 2
    missing_subtitles: list[str] = []
    for idx, style, text in reader_headings:
        if style != "Heading 1":
            continue
        if text.startswith("END OF BOOK"):
            continue
        # find next heading
        nxt = None
        for j, s2, t2 in reader_headings:
            if j > idx:
                nxt = (j, s2, t2)
                break
        if not nxt or nxt[1] != "Heading 2":
            missing_subtitles.append(text)

    flags = []
    for pat in FLAG_PATTERNS:
        if re.search(pat, reader_text, re.I):
            flags.append(f"Reader contains flagged token pattern: `{pat}`")

    if cfg.end_marker not in reader_text:
        flags.append(f"Reader missing end marker: {cfg.end_marker}")

    # Compare reader/master deltas at a high level
    reader_wc = words_count(reader_text)
    master_wc = words_count(master_text)
    delta = abs(master_wc - reader_wc)

    stats = {
        "reader_words": reader_wc,
        "master_words": master_wc,
        "word_delta_reader_master": delta,
        "reader_heading1_count": len(h1_actual),
        "reader_heading2_count": len([1 for _, s, _ in reader_headings if s == "Heading 2"]),
        "missing_heading1_expected": len(missing),
        "missing_subtitles": len(missing_subtitles),
        "flag_token_matches": len(flags),
    }

    status = "PASS" if not missing and not missing_subtitles and not flags else "FLAGGED"

    lines = [
        f"# Publication Prep Audit — {cfg.title}",
        "",
        f"- Status: **{status}**",
        f"- Reader: `{reader_path}`",
        f"- Master: `{master_path}`",
        "",
        "## Stats",
    ]
    for k, v in stats.items():
        lines.append(f"- {k}: {v}")

    lines.append("")
    lines.append("## Heading checks")
    if missing:
        lines.append("- Missing Heading 1 entries:")
        lines.extend([f"  - {m}" for m in missing])
    else:
        lines.append("- Heading 1 sequence present")

    if missing_subtitles:
        lines.append("- Missing chapter/prologue/interlude subtitles:")
        lines.extend([f"  - {m}" for m in missing_subtitles])
    else:
        lines.append("- Subtitle structure present")

    lines.append("")
    lines.append("## Flagged token checks")
    if flags:
        lines.extend([f"- {f}" for f in flags])
    else:
        lines.append("- No flagged placeholders/tokens in reader copy")

    lines.append("")
    lines.append("## First non-empty lines (reader)")
    first = [p.text.strip() for p in reader_doc.paragraphs if p.text.strip()][:8]
    for i, ln in enumerate(first, 1):
        lines.append(f"{i}. {ln}")

    lines.append("")
    lines.append("## Final non-empty lines (reader)")
    tail = [p.text.strip() for p in reader_doc.paragraphs if p.text.strip()][-8:]
    for i, ln in enumerate(tail, 1):
        lines.append(f"{i}. {ln}")

    return "\n".join(lines) + "\n", stats, ([] if status == "PASS" else (missing + missing_subtitles + flags))


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    summary_rows = []
    all_flags = 0
    for cfg in BOOKS:
        body, stats, flags = audit_book(cfg)
        out_path = OUT_DIR / f"{cfg.key}-audit.md"
        out_path.write_text(body, encoding="utf-8")
        all_flags += len(flags)
        summary_rows.append((cfg.title, stats, len(flags), out_path.name))
        print(f"{cfg.key}: flags={len(flags)} words={stats['reader_words']}")

    lines = [
        "# Final Publication-Prep Pass — Trilogy QA Summary",
        "",
        f"Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
        "",
        "| Book | Reader words | Heading gaps | Subtitle gaps | Token flags | Status | Report |",
        "|------|--------------|--------------|---------------|------------:|--------|--------|",
    ]
    for title, stats, flag_count, report in summary_rows:
        status = "PASS" if flag_count == 0 else "FLAGGED"
        lines.append(
            f"| {title} | {stats['reader_words']} | {stats['missing_heading1_expected']} | {stats['missing_subtitles']} | {stats['flag_token_matches']} | {status} | `{report}` |"
        )

    lines.extend(
        [
            "",
            f"Overall: **{'PASS' if all_flags == 0 else 'FLAGGED'}**",
            "",
            "Reports are in `production/publication-prep/audit/`.",
        ]
    )
    (OUT_DIR / "SUMMARY.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
