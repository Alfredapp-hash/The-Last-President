#!/usr/bin/env python3
"""Random chapter spot-check for author notes before Gumroad upload."""

from __future__ import annotations

import argparse
import random
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from docx import Document


ROOT = Path("/workspace")
BOOKS_DIR = ROOT / "production" / "books"
OUT_DIR = ROOT / "production" / "publication-prep" / "gumroad-spot-check"

AUTHOR_NOTE_PATTERNS = [
    (r"\[ILLUSTRATION PLACEMENT\]", "illustration placement"),
    (r"\bAUTHOR(?:'S)?\s+NOTE\b", "author note"),
    (r"\bEDITOR(?:'S)?\s+NOTE\b", "editor note"),
    (r"\bPRODUCTION\s+NOTE\b", "production note"),
    (r"\bTODO\b", "TODO"),
    (r"\bFIXME\b", "FIXME"),
    (r"\bTK\b", "TK"),
    (r"\bTBD\b", "TBD"),
    (r"\bPLACEHOLDER\b", "PLACEHOLDER"),
    (r"\[NOTE[:\s]", "bracketed NOTE"),
    (r"\[EDIT[:\s]", "bracketed EDIT"),
    (r"<!--", "HTML comment"),
    (r"\bLorem ipsum\b", "lorem ipsum"),
    (r"\bXXX\b", "XXX marker"),
    (r"\?\?\?", "??? marker"),
    (r"\[ART\b", "ART bracket"),
    (r"\bART\s+BRIEF\b", "ART BRIEF"),
    (r"\bQUERY:\b", "QUERY"),
    (r"\bSTET\b", "STET"),
    (r"\[CHAPTER\s+STUB\]", "chapter stub"),
    (r"\bNEEDS\s+(?:REWRITE|EDIT|WORK)\b", "needs work"),
    (r"\bCHECK\s+WITH\s+AUTHOR\b", "check with author"),
    (r"\[internal\]", "internal bracket"),
]

META_NOTE_PATTERNS = [
    (r"You're around Chapter", "chapter-progress author note"),
    (r"I'd estimate \d+[-–]\d+ chapters left", "chapter-count author note"),
    (r"remaining arc should probably be", "outline author note"),
    (r"late Act II\s*/\s*early Act III", "structure author note"),
    (r"For this satirical thriller, I'd treat", "genre author note"),
]

KNOWN_AUTHOR_NOTE_PATTERNS = [
    re.compile(r"You.re around Chapter \d+\.", re.I),
    re.compile(r"For this satirical thriller, I.d treat this as the late Act II", re.I),
    re.compile(
        r"I.d estimate \d+[-–]\d+ chapters left if we want a full, satisfying ending",
        re.I,
    ),
    re.compile(r"The remaining arc should probably be:", re.I),
    re.compile(
        r"Old Court exposure → Dahlia cornered → adult target names begin unsealing",
        re.I,
    ),
]


@dataclass
class BookConfig:
    key: str
    title: str
    reader_docx: str
    master_docx: str


BOOKS = [
    BookConfig(
        key="book-one",
        title="The Last President",
        reader_docx="Baren_Sump_BOOK_ONE_The_Last_President_FINAL_READER_COPY.docx",
        master_docx="Baren_Sump_BOOK_ONE_The_Last_President_AUTHOR_PRODUCTION_MASTER.docx",
    ),
    BookConfig(
        key="book-two",
        title="Children of Tomorrow",
        reader_docx="Baren_Sump_BOOK_TWO_Children_of_Tomorrow_FINAL_READER_COPY.docx",
        master_docx="Baren_Sump_BOOK_TWO_Children_of_Tomorrow_AUTHOR_PRODUCTION_MASTER.docx",
    ),
    BookConfig(
        key="book-three",
        title="The Black Path",
        reader_docx="Baren_Sump_BOOK_THREE_The_Black_Path_FINAL_READER_COPY.docx",
        master_docx="Baren_Sump_BOOK_THREE_The_Black_Path_AUTHOR_PRODUCTION_MASTER.docx",
    ),
]


def words_count(text: str) -> int:
    return len(re.findall(r"\b[\w'’-]+\b", text))


def split_sections(doc: Document) -> list[dict]:
    sections: list[dict] = []
    current: dict | None = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and text:
            if current:
                sections.append(current)
            current = {"heading": text, "lines": []}
        elif current is not None and text:
            current["lines"].append((i + 1, text))
    if current:
        sections.append(current)
    return sections


def scan_text(text: str, line_lookup: list[tuple[int, str]]) -> list[dict]:
    issues: list[dict] = []
    all_patterns = AUTHOR_NOTE_PATTERNS + META_NOTE_PATTERNS
    for pat, label in all_patterns:
        for m in re.finditer(pat, text, re.I | re.M):
            line_no = None
            context = ""
            for ln, txt in line_lookup:
                if m.group(0) in txt:
                    line_no = ln
                    context = txt
                    break
            if line_no is None:
                context = text[max(0, m.start() - 40) : m.end() + 40]
            issues.append(
                {
                    "label": label,
                    "token": m.group(0),
                    "line": line_no,
                    "context": context,
                }
            )
    return issues


def remove_known_author_notes(path: Path) -> list[str]:
    doc = Document(str(path))
    removed: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if any(pat.search(text) for pat in KNOWN_AUTHOR_NOTE_PATTERNS):
            removed.append(text)
            p._element.getparent().remove(p._element)
    if removed:
        doc.save(str(path))
    return removed


def spot_check_book(cfg: BookConfig, seed: int, samples: int = 3) -> tuple[str, list[dict]]:
    path = BOOKS_DIR / cfg.key / cfg.reader_docx
    doc = Document(str(path))
    sections = split_sections(doc)
    pool = [s for s in sections if not s["heading"].startswith("END OF BOOK")]
    rng = random.Random(seed + hash(cfg.key) % 1000)
    chosen = rng.sample(pool, min(samples, len(pool)))

    lines: list[str] = [
        f"## {cfg.title}",
        "",
        f"- Source: `{path}`",
        f"- Sampled sections: {', '.join(s['heading'] for s in chosen)}",
        "",
    ]
    book_issues: list[dict] = []

    for sec in chosen:
        body = "\n".join(t for _, t in sec["lines"])
        wc = words_count(body)
        issues = scan_text(body, sec["lines"])
        book_issues.extend({**i, "section": sec["heading"]} for i in issues)
        lines.append(f"### {sec['heading']} ({wc:,} words)")
        if issues:
            lines.append("- **FLAGS:**")
            for issue in issues:
                lines.append(
                    f"  - `{issue['label']}` line {issue['line']}: `{issue['token']}`"
                )
                lines.append(f"    - {issue['context'][:220]}")
        else:
            lines.append("- No author-note patterns detected.")
        preview = sec["lines"][:2] + sec["lines"][-2:]
        lines.append("- Sample lines:")
        seen: set[int] = set()
        for ln, txt in preview:
            if ln in seen:
                continue
            seen.add(ln)
            lines.append(f"  - L{ln}: {txt[:180]}{'...' if len(txt) > 180 else ''}")
        lines.append("")

    return "\n".join(lines), book_issues


def full_manuscript_meta_scan(cfg: BookConfig) -> list[dict]:
    path = BOOKS_DIR / cfg.key / cfg.reader_docx
    doc = Document(str(path))
    rows: list[tuple[int, str, str]] = []
    current_h1 = ""
    for i, p in enumerate(doc.paragraphs, 1):
        text = p.text.strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and text:
            current_h1 = text
        if text:
            rows.append((i, current_h1, text))
    hits: list[dict] = []
    for ln, h1, text in rows:
        for pat, label in META_NOTE_PATTERNS:
            if re.search(pat, text, re.I):
                hits.append({"line": ln, "section": h1, "label": label, "context": text})
                break
    return hits


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--seed", type=int, default=20260622)
    parser.add_argument("--samples", type=int, default=3)
    parser.add_argument("--fix-known-notes", action="store_true")
    args = parser.parse_args()

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    report = [
        "# Gumroad Random Chapter Spot Check",
        "",
        f"Date: {now}",
        f"Seed: {args.seed}",
        f"Samples per book: {args.samples}",
        "",
    ]

    if args.fix_known_notes:
        report.append("## Remediation")
        for cfg in BOOKS:
            for kind in (cfg.reader_docx, cfg.master_docx):
                path = BOOKS_DIR / cfg.key / kind
                removed = remove_known_author_notes(path)
                if removed:
                    report.append(f"- Removed {len(removed)} author-note paragraph(s) from `{path}`")
                    for item in removed:
                        report.append(f"  - {item[:180]}{'...' if len(item) > 180 else ''}")
        report.append("")

    all_issues: list[dict] = []
    for cfg in BOOKS:
        body, issues = spot_check_book(cfg, seed=args.seed, samples=args.samples)
        report.append(body)
        all_issues.extend(issues)
        meta_hits = full_manuscript_meta_scan(cfg)
        report.append("### Full-manuscript meta-note scan")
        if meta_hits:
            report.append("- **FLAGS:**")
            for hit in meta_hits:
                report.append(
                    f"  - L{hit['line']} [{hit['section']}] `{hit['label']}`: {hit['context'][:220]}"
                )
            all_issues.extend(meta_hits)
        else:
            report.append("- No meta-note patterns detected.")
        report.append("")

    status = "PASS" if not all_issues else "FLAGGED"
    report.insert(4, f"Overall: **{status}**")
    report.insert(5, "")

    out_path = OUT_DIR / "SPOT_CHECK_REPORT.md"
    out_path.write_text("\n".join(report) + "\n", encoding="utf-8")
    print(f"Wrote {out_path}")
    print(f"Overall: {status} ({len(all_issues)} flags)")


if __name__ == "__main__":
    main()
