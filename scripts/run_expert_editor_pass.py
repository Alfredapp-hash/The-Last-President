#!/usr/bin/env python3
"""
Expert Editor Pass Runner — Baren Sump Trilogy
Implements Series Master Phases 1-4 per production/editor-prompts/
"""

from __future__ import annotations

import re
import shutil
import zipfile
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

BOOKS = Path("/workspace/production/books")

WT = re.compile(r"(<w:t(?: xml:space=\"preserve\")?>)([^<]*?)(</w:t>)", re.S)

# Intentional Barren (wordplay) — exclude from typo fixes
BAREN_WORDPLAY_CTX = re.compile(
    r"BARREN\s*/\s*BAREN|Barren\?|barren\?|misspelled his name|variant of barren",
    re.I,
)

NAME_VARIANTS = {
    "Chronold": "Chronald",
    "Moro": "Morro",
    "Veira": "Veyra",
    "Dalia": "Dahlia",
    "Lilah": "Lila",
}

CONTROLLING = {
    "book-one": [
        "Every accusation is advertising",
        "To him not knowing until morning",
    ],
    "book-two": [
        "Children are not evidence",
        "Some days were allowed to end without becoming part of the case",
        "He gets to choose his own face",
    ],
    "book-three": [
        "The black path was not a tunnel beneath the house",
        "Some days refused to become evidence",
        "Not as proof. Not as witness. Not as remedy",
        "And all, for once, did not ask to be everything",
    ],
}

SPACING_DEFECTS = [
    "writtenChildren",
    "calledThe",
    "circledlaughed",
    "andstabilization",
    "askedwell",
    "readTHE",
    "AtBAREN",
    "triedTHELASTPRESIDENT",
    "underlinedfunction",
    "outmay",
]

READER_MARKERS = re.compile(
    r"\[ILLUSTRATION PLACEMENT\]|\bTK\b|\bTODO\b|\bFIXME\b|\bPLACEHOLDER\b",
    re.I,
)


@dataclass
class EditStats:
    run_boundary_fixes: int = 0
    in_run_fixes: int = 0
    double_space_fixes: int = 0
    dash_fixes: int = 0
    trailing_space_fixes: int = 0
    marker_removals: int = 0
    name_fixes: int = 0
    changes: list[str] = field(default_factory=list)
    flags: list[str] = field(default_factory=list)


def needs_space_between(t1: str, t2: str) -> bool:
    if not t1 or not t2 or t2[0] == " " or t1[-1] == " " or t1[-1] == "-":
        return False
    return t1[-1].isalnum() and t2[0].isalnum()


def fix_wt_content(text: str, stats: EditStats, is_reader: bool) -> str:
    orig = text
    text = re.sub(r"  +", " ", text)
    if text != orig and "  " in orig:
        stats.double_space_fixes += 1

    orig = text
    text = re.sub(r"([^\s])--([^\s])", r"\1—\2", text)
    text = re.sub(r" -- ", " — ", text)
    if text != orig:
        stats.dash_fixes += 1

    orig = text
    text = text.rstrip(" ")
    if text != orig:
        stats.trailing_space_fixes += 1

    if is_reader:
        orig = text
        text = READER_MARKERS.sub("", text)
        if text != orig:
            stats.marker_removals += 1

    for bad, good in NAME_VARIANTS.items():
        if bad in text and not BAREN_WORDPLAY_CTX.search(text):
            text = text.replace(bad, good)
            stats.name_fixes += 1
            stats.changes.append(f"name: {bad} → {good}")

    subs = [
        (r"\b(tried|read|said|used|At)([A-Z]{2,})", r"\1 \2"),
        (r"\b(circled)(laughed)\b", r"\1 \2"),
        (r"\b(underlined)(function)\b", r"\1 \2"),
        (r"\b(crossed out)(may)\b", r"\1 \2"),
        (r"\b(asked)(well)\b", r"\1 \2"),
        (r"\b(and)(stabilization)\b", r"\1 \2"),
        (r"THELASTPRESIDENT", "THE LAST PRESIDENT"),
    ]
    for pat, repl in subs:
        new, n = re.subn(pat, repl, text)
        if n:
            stats.in_run_fixes += n
            text = new

    return text


def fix_paragraph_xml(p_xml: str, stats: EditStats, is_reader: bool) -> str:
    def fix_text(m: re.Match) -> str:
        o, c, cl = m.group(1), m.group(2), m.group(3)
        nc = fix_wt_content(c, stats, is_reader)
        return o + nc + cl

    p_xml = WT.sub(fix_text, p_xml)
    parts = WT.split(p_xml)
    if len(parts) >= 4:
        rebuilt = [parts[0]]
        i, prev = 1, None
        while i + 2 < len(parts):
            ot, tx, ct = parts[i], parts[i + 1], parts[i + 2]
            if prev is not None and needs_space_between(prev, tx) and not tx.startswith(" "):
                tx = " " + tx
                ot = '<w:t xml:space="preserve">'
                stats.run_boundary_fixes += 1
            rebuilt.extend([ot, tx, ct])
            prev = tx
            i += 3
            if i < len(parts):
                rebuilt.append(parts[i])
                i += 1
        p_xml = "".join(rebuilt)
    return p_xml


def patch_docx(src: Path, dst: Path, is_reader: bool, stats: EditStats) -> None:
    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                xml = data.decode("utf-8")
                paragraphs = re.split(r"(</w:p>)", xml)
                out = []
                for i in range(0, len(paragraphs), 2):
                    chunk = paragraphs[i]
                    closer = paragraphs[i + 1] if i + 1 < len(paragraphs) else ""
                    if "<w:p" in chunk:
                        chunk = fix_paragraph_xml(chunk + closer, stats, is_reader)
                        out.append(chunk)
                    else:
                        out.append(chunk + closer)
                data = "".join(out).encode("utf-8")
            zout.writestr(item, data)


def extract_plain(docx: Path) -> str:
    with zipfile.ZipFile(docx) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    return re.sub(r"<[^>]+>", "", xml)


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w'\u2019-]+\b", text))


def verify_book(book_key: str, reader: Path, stats: EditStats) -> dict:
    plain = extract_plain(reader)
    results = {}

    for phrase in CONTROLLING[book_key]:
        ok = phrase.lower() in plain.lower()
        results[f"controlling:{phrase[:40]}"] = ok
        if not ok:
            stats.flags.append(f"MISSING controlling line: {phrase}")

    for defect in SPACING_DEFECTS:
        if defect in plain:
            stats.flags.append(f"SPACING DEFECT remains: {defect}")
            results[f"defect:{defect}"] = False
        else:
            results[f"defect:{defect}"] = True

    for variant in NAME_VARIANTS:
        if variant in plain:
            # Barren exception only book one
            if variant == "Barren":
                continue
            stats.flags.append(f"Name variant found: {variant}")
            results[f"name:{variant}"] = False

    if book_key == "book-one" and "Barren" in plain:
        results["barren_wordplay_present"] = True

    if READER_MARKERS.search(plain):
        stats.flags.append("Production markers remain in reader copy")
        results["markers_clean"] = False
    else:
        results["markers_clean"] = True

    results["word_count"] = word_count(plain)
    results["double_space_runs"] = len(re.findall(r"  +", plain))
    return results


def verify_headings(docx: Path, book_key: str, stats: EditStats) -> None:
    try:
        from docx import Document
    except ImportError:
        stats.flags.append("python-docx unavailable for heading check")
        return

    doc = Document(str(docx))
    headings = [
        (p.style.name if p.style else "", p.text.strip())
        for p in doc.paragraphs
        if p.text.strip() and (p.style.name.startswith("Heading") if p.style else False)
    ]

    if book_key == "book-one":
        h1 = [t for s, t in headings if s.startswith("Heading 1")]
        if "INTERLUDE" not in h1:
            stats.flags.append("Book One INTERLUDE not Heading 1")
        if "END OF BOOK ONE" not in " ".join(h1):
            if not any("END OF BOOK ONE" in t for s, t in headings):
                stats.flags.append("Missing END OF BOOK ONE heading")

    if book_key == "book-two":
        h2 = [t for s, t in headings if s.startswith("Heading 2")]
        if "They Carried Moral Force" not in h2:
            stats.flags.append("Book Two Ch.30 subtitle missing")

    if book_key == "book-three":
        plain_titles = sum(
            1
            for p in doc.paragraphs[:15]
            if p.text.strip() == "The Black Path" and not (p.style and p.style.name.startswith("Heading"))
        )
        if plain_titles > 1:
            stats.flags.append("Book Three duplicate title lines on title page")


def run_book(book_key: str, reader_name: str, master_name: str) -> EditStats:
    book_dir = BOOKS / book_key
    reader = book_dir / reader_name
    master = book_dir / master_name
    stats = EditStats()

    for path, is_reader in [(reader, True), (master, False)]:
        tmp = path.with_suffix(".tmp.docx")
        patch_docx(path, tmp, is_reader, stats)
        shutil.move(tmp, path)

    verify_headings(reader, book_key, stats)
    verification = verify_book(book_key, reader, stats)

    floors = {"book-one": 82500, "book-two": 87400, "book-three": 61000}
    wc = verification["word_count"]
    if wc < floors[book_key]:
        stats.flags.append(f"Word count {wc} below floor {floors[book_key]}")

    log_path = book_dir / "PRODUCTION_EDITOR_LOG.txt"
    lines = [
        f"{book_key.upper()} — Expert Editor Pass",
        f"Date: {date.today().isoformat()}",
        f"Agent: run_expert_editor_pass.py (Series Master Phases 1-4)",
        "",
        "## Edit counts",
        f"- Run-boundary spacing fixes: {stats.run_boundary_fixes}",
        f"- In-run text fixes: {stats.in_run_fixes}",
        f"- Double-space fixes: {stats.double_space_fixes}",
        f"- Dash normalizations: {stats.dash_fixes}",
        f"- Trailing space fixes: {stats.trailing_space_fixes}",
        f"- Production marker removals: {stats.marker_removals}",
        f"- Name variant fixes: {stats.name_fixes}",
        "",
        f"## Word count: {wc:,}",
        "",
        "## Verification",
    ]
    for k, v in sorted(verification.items()):
        lines.append(f"- {k}: {v}")
    if stats.changes:
        lines.extend(["", "## Sample changes"] + [f"- {c}" for c in stats.changes[:20]])
    if stats.flags:
        lines.extend(["", "## FLAGS (author review)"] + [f"- {f}" for f in stats.flags])
    else:
        lines.extend(["", "## Status: PASS — no flags"])
    log_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return stats


BOOK_CONFIG = [
    (
        "book-one",
        "Baren_Sump_BOOK_ONE_The_Last_President_FINAL_READER_COPY.docx",
        "Baren_Sump_BOOK_ONE_The_Last_President_AUTHOR_PRODUCTION_MASTER.docx",
    ),
    (
        "book-two",
        "Baren_Sump_BOOK_TWO_Children_of_Tomorrow_FINAL_READER_COPY.docx",
        "Baren_Sump_BOOK_TWO_Children_of_Tomorrow_AUTHOR_PRODUCTION_MASTER.docx",
    ),
    (
        "book-three",
        "Baren_Sump_BOOK_THREE_The_Black_Path_FINAL_READER_COPY.docx",
        "Baren_Sump_BOOK_THREE_The_Black_Path_AUTHOR_PRODUCTION_MASTER.docx",
    ),
]


def write_checklist(all_stats: dict[str, EditStats]) -> None:
    out = BOOKS.parent / "editor-prompts" / "PRODUCTION_READINESS_SIGNED.txt"
    total_flags = sum(len(s.flags) for s in all_stats.values())
    status = "READY FOR QUERY/BETA" if total_flags == 0 else "NOT READY — see flags"
    lines = [
        "PRODUCTION READINESS — AUTOMATED SIGN-OFF",
        f"Date: {date.today().isoformat()}",
        f"Series status: {status}",
        "",
    ]
    for key, st in all_stats.items():
        lines.append(f"### {key}")
        lines.append(f"  Flags: {len(st.flags)}")
        for f in st.flags:
            lines.append(f"    - {f}")
        lines.append("")
    out.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    print("=== SERIES MASTER (Phases 1-4) — applying to all books ===\n")
    all_stats = {}
    for book_key, reader, master in BOOK_CONFIG:
        print(f"--- {book_key} ---")
        st = run_book(book_key, reader, master)
        all_stats[book_key] = st
        print(f"  run_boundary={st.run_boundary_fixes} in_run={st.in_run_fixes} flags={len(st.flags)}")
    write_checklist(all_stats)
    print("\n=== DONE ===")
    total_flags = sum(len(s.flags) for s in all_stats.values())
    print(f"Total flags: {total_flags}")


if __name__ == "__main__":
    main()
