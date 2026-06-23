#!/usr/bin/env python3
"""Build professional opening-chapter reader samples (EPUB + PDF with cover)."""

from __future__ import annotations

import html
import os
import shutil
import subprocess
from dataclasses import dataclass, replace
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "public" / "downloads"
WORK_DIR = ROOT / "production" / "sample-assets" / "build"
PDF_CSS = ROOT / "production" / "sample-assets" / "reader-sample.css"
EPUB_CSS = ROOT / "production" / "sample-assets" / "epub-sample.css"
SERIES_TITLE = "Baren Sump and The Last President"
AUTHOR = "William Sailsbury"
CONTACT = "hello@thesumpledger.com"
WEBSITE = "https://www.thesumpledger.com"

CHROME_PATHS = [
    "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
    "/Applications/Chromium.app/Contents/MacOS/Chromium",
]


@dataclass(frozen=True)
class BookSampleConfig:
    slug: str
    title: str
    subtitle: str
    book_number: int
    accent: str
    gumroad_url: str
    source_docx: Path
    cover_image: Path


BOOKS = [
    BookSampleConfig(
        slug="the-last-president",
        title="The Last President",
        subtitle="Book One",
        book_number=1,
        accent="#9e2b3c",
        gumroad_url="https://salsbury61.gumroad.com/l/lneqtc",
        source_docx=ROOT
        / "production/books/book-one/Baren_Sump_BOOK_ONE_The_Last_President_FINAL_READER_COPY.docx",
        cover_image=ROOT / "public/images/covers/cover-book-one-titled.png",
    ),
    BookSampleConfig(
        slug="children-of-tomorrow",
        title="Children of Tomorrow",
        subtitle="Book Two",
        book_number=2,
        accent="#c9a962",
        gumroad_url="https://salsbury61.gumroad.com/l/qxixz",
        source_docx=ROOT
        / "production/books/book-two/Baren_Sump_BOOK_TWO_Children_of_Tomorrow_FINAL_READER_COPY.docx",
        cover_image=ROOT / "public/images/covers/cover-book-two-titled.png",
    ),
    BookSampleConfig(
        slug="the-black-path",
        title="The Black Path",
        subtitle="Book Three",
        book_number=3,
        accent="#5a7a52",
        gumroad_url="https://salsbury61.gumroad.com/l/hqtvrx",
        source_docx=ROOT
        / "production/books/book-three/Baren_Sump_BOOK_THREE_The_Black_Path_FINAL_READER_COPY.docx",
        cover_image=ROOT / "public/images/covers/cover-book-three-titled.png",
    ),
]

ORDINALS = [
    "ONE",
    "TWO",
    "THREE",
    "FOUR",
    "FIVE",
    "SIX",
    "SEVEN",
    "EIGHT",
    "NINE",
    "TEN",
]


@dataclass
class Paragraph:
    kind: str
    text: str


@dataclass
class Section:
    heading: str
    display_heading: str
    paragraphs: list[Paragraph]


def is_narrative_heading(text: str) -> bool:
    t = text.strip().upper()
    return t.startswith("CHAPTER ") or t in {"PROLOGUE", "INTERLUDE", "THE BLACK PATH"}


def paragraph_kind(style_name: str) -> str:
    style = (style_name or "").lower()
    if style == "heading 2":
        return "h2"
    if style == "heading 1":
        return "h1"
    return "p"


def extract_sections(doc: Document) -> list[Section]:
    paras = doc.paragraphs
    heading_idxs: list[tuple[int, str]] = []
    for i, p in enumerate(paras):
        style = (p.style.name if p.style else "") or ""
        txt = p.text.strip()
        if not txt:
            continue
        if style == "Heading 1" and is_narrative_heading(txt):
            heading_idxs.append((i, txt))

    sections: list[Section] = []
    for n, (start_idx, heading) in enumerate(heading_idxs):
        end_idx = heading_idxs[n + 1][0] if n + 1 < len(heading_idxs) else len(paras)
        items: list[Paragraph] = []
        for p in paras[start_idx + 1 : end_idx]:
            txt = p.text.strip()
            if not txt:
                continue
            items.append(Paragraph(paragraph_kind(p.style.name if p.style else ""), txt))
        sections.append(
            Section(heading=heading, display_heading=heading, paragraphs=items)
        )
    return sections


def select_opening_sample(sections: list[Section]) -> list[Section]:
    """Prologue (Book One only) + first three chapter sections of this volume."""
    selected: list[Section] = []
    start = 0
    if sections and sections[0].heading == "PROLOGUE":
        selected.append(sections[0])
        start = 1

    chapter_sections = [s for s in sections[start:] if s.heading.startswith("CHAPTER ")]
    selected.extend(chapter_sections[:3])
    return selected


def relabel_for_sample(sections: list[Section]) -> list[Section]:
    """Use reader-facing chapter labels in the sample (Ch 1–3) while keeping content."""
    relabeled: list[Section] = []
    chapter_idx = 0
    for section in sections:
        if section.heading == "PROLOGUE":
            relabeled.append(section)
            continue
        if section.heading.startswith("CHAPTER "):
            chapter_idx += 1
            label = f"CHAPTER {ORDINALS[chapter_idx - 1]}"
            relabeled.append(
                replace(section, display_heading=label)
            )
        else:
            relabeled.append(section)
    return relabeled


def scope_line(cfg: BookSampleConfig, sections: list[Section]) -> str:
    includes_prologue = bool(sections and sections[0].heading == "PROLOGUE")
    chapter_headings = [s.heading for s in sections if s.heading.startswith("CHAPTER ")]

    if cfg.book_number == 1:
        return (
            "This complimentary reader sample includes the prologue and the opening "
            "three chapters of Book One."
        )

    trilogy_nums = ", ".join(chapter_headings)
    prerequisite = "Read Book One before starting Book Two." if cfg.book_number == 2 else (
        "Read Books One and Two before starting Book Three."
    )
    return (
        f"This complimentary reader sample includes the opening three chapters of "
        f"{cfg.subtitle}. {prerequisite} "
        f"(Trilogy manuscript numbering: {trilogy_nums}.)"
    )


def render_markdown(cfg: BookSampleConfig, sections: list[Section]) -> str:
    lines = [
        "---",
        f'title: "{cfg.title}"',
        f'subtitle: "{cfg.subtitle} — Opening Chapters Reader Sample"',
        f'creator: "{AUTHOR}"',
        f"rights: © {AUTHOR}. Sample excerpt for reader evaluation.",
        "---",
        "",
        f"% {cfg.title}",
        "",
        f"*{cfg.subtitle} · {SERIES_TITLE}*",
        "",
        "## Reader Sample",
        "",
        scope_line(cfg, sections),
        "",
        f"Full volume: {WEBSITE}/books/{cfg.slug}",
        f"Buy: {cfg.gumroad_url}",
        f"Rights, review copies, and media: {CONTACT}",
        "",
        "---",
        "",
    ]

    for section in sections:
        lines.append(f"# {section.display_heading}")
        lines.append("")
        for para in section.paragraphs:
            if para.kind in {"h2", "h1"}:
                lines.append(f"## {para.text}")
                lines.append("")
            else:
                lines.append(para.text)
                lines.append("")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def render_html(cfg: BookSampleConfig, sections: list[Section]) -> str:
    chapter_blocks: list[str] = []
    for section in sections:
        body: list[str] = []
        for para in section.paragraphs:
            safe = html.escape(para.text)
            if para.kind in {"h2", "h1"}:
                body.append(f'<h3 class="chapter-title">{safe}</h3>')
            else:
                body.append(f"<p>{safe}</p>")
        chapter_blocks.append(
            f'<section class="chapter">'
            f'<h2 class="chapter-label">{html.escape(section.display_heading)}</h2>'
            f"{''.join(body)}"
            f"</section>"
        )

    cover_uri = cfg.cover_image.resolve().as_uri()
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>{html.escape(cfg.title)} — Reader Sample</title>
  <link rel="stylesheet" href="{PDF_CSS.resolve().as_uri()}" />
  <style>:root {{ --accent: {cfg.accent}; }}</style>
</head>
<body>
  <section class="cover-page">
    <div class="cover-frame">
      <img src="{cover_uri}" alt="{html.escape(cfg.title)} cover" />
    </div>
    <p class="cover-badge">Reader Sample · Opening Chapters</p>
    <p class="cover-series">{html.escape(SERIES_TITLE)}</p>
  </section>

  <section class="front-matter">
    <p class="eyebrow">{html.escape(SERIES_TITLE)}</p>
    <h1 class="title">{html.escape(cfg.title)}</h1>
    <p class="subtitle">{html.escape(cfg.subtitle)}</p>
    <div class="sample-box">
      <h2>Complimentary excerpt</h2>
      <p>{html.escape(scope_line(cfg, sections))}</p>
      <p>Buy the full volume: <a href="{html.escape(cfg.gumroad_url)}">{html.escape(cfg.gumroad_url)}</a></p>
      <p>Series site: {WEBSITE}/books/{html.escape(cfg.slug)} · Review copies: {CONTACT}</p>
    </div>
    <p class="meta-line">{html.escape(AUTHOR)} · Reader Sample Edition</p>
  </section>

  {''.join(chapter_blocks)}

  <section class="back-page">
    <h2>Continue reading</h2>
    <p>The complete volume is available now on Gumroad.</p>
    <p class="back-url">{html.escape(cfg.gumroad_url)}</p>
    <p class="back-url">{WEBSITE}/books/{html.escape(cfg.slug)}</p>
    <p class="meta-line">© {html.escape(AUTHOR)} · {html.escape(SERIES_TITLE)}</p>
  </section>
</body>
</html>
"""


def run(cmd: list[str]) -> None:
    result = subprocess.run(cmd, text=True, capture_output=True)
    if result.returncode != 0:
        raise RuntimeError(
            f"Command failed ({result.returncode}): {' '.join(cmd)}\n"
            f"{result.stderr or result.stdout}"
        )


def html_to_pdf(html_path: Path, pdf_path: Path) -> None:
    if shutil.which("wkhtmltopdf"):
        run(
            [
                "wkhtmltopdf",
                "--enable-local-file-access",
                "--margin-top",
                "0",
                "--margin-bottom",
                "0",
                "--margin-left",
                "0",
                "--margin-right",
                "0",
                str(html_path),
                str(pdf_path),
            ]
        )
        return

    chrome = next((p for p in CHROME_PATHS if Path(p).exists()), None)
    if not chrome:
        raise RuntimeError("No PDF renderer found (wkhtmltopdf or Chrome).")

    run(
        [
            chrome,
            "--headless=new",
            "--disable-gpu",
            "--no-pdf-header-footer",
            f"--print-to-pdf={pdf_path}",
            html_path.resolve().as_uri(),
        ]
    )


def build_sample(cfg: BookSampleConfig) -> None:
    doc = Document(str(cfg.source_docx))
    raw = select_opening_sample(extract_sections(doc))
    if not raw:
        raise RuntimeError(f"No narrative sections found in {cfg.source_docx}")
    sections = relabel_for_sample(raw)

    WORK_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    base = cfg.slug
    md_path = WORK_DIR / f"{base}.md"
    html_path = WORK_DIR / f"{base}.html"
    epub_path = OUT_DIR / f"{base}-first-3-chapters-sample.epub"
    pdf_path = OUT_DIR / f"{base}-first-3-chapters-sample.pdf"

    md_path.write_text(render_markdown(cfg, sections), encoding="utf-8")
    html_path.write_text(render_html(cfg, sections), encoding="utf-8")

    run(
        [
            "pandoc",
            str(md_path),
            "--from=markdown",
            "--to=epub3",
            "--css",
            str(EPUB_CSS),
            "--epub-cover-image",
            str(cfg.cover_image),
            "--metadata",
            f"title={cfg.title}",
            "--metadata",
            f"subtitle={cfg.subtitle} — Opening Chapters Reader Sample",
            "--metadata",
            f"creator={AUTHOR}",
            "--metadata",
            f"description={scope_line(cfg, sections)}",
            "-o",
            str(epub_path),
        ]
    )

    html_to_pdf(html_path, pdf_path)

    chapter_count = sum(1 for s in sections if s.heading.startswith("CHAPTER "))
    prologue = sections[0].heading == "PROLOGUE" if sections else False
    labels = [s.display_heading for s in sections]
    print(
        f"Wrote {epub_path.name} and {pdf_path.name} "
        f"({chapter_count} chapters{' + prologue' if prologue else ''}: {', '.join(labels)})"
    )


def main() -> None:
    for cfg in BOOKS:
        build_sample(cfg)


if __name__ == "__main__":
    main()